import os
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 153)

doc = Document()

# Tiêu đề
title = doc.add_heading('BÁO CÁO TÍNH NĂNG & LỘ TRÌNH PHÁT TRIỂN\nHỆ THỐNG AI INVOICE', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('==================================================================').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('\n')

# Phần 1: CÁC TÍNH NĂNG ĐÃ CÓ
add_heading(doc, 'PHẦN 1: CÁC TÍNH NĂNG ĐÃ CÓ TRONG HỆ THỐNG (HIỆN TẠI)')
p1 = doc.add_paragraph('Dự án đã hoàn thiện được lõi (Core Engine) cơ bản, cụ thể bao gồm:\n')

doc.add_paragraph('1. Trình giả lập dữ liệu (Mock Data Generator):', style='List Bullet')
p1_1 = doc.add_paragraph()
p1_1.add_run('Đã có script (generate_mock_invoices.py) tự động sinh ra hàng loạt hóa đơn giả lập (MISA, FAST, DEFAULT) với độ nhiễu màu sắc, font chữ ngẫu nhiên để huấn luyện và kiểm thử AI.')

doc.add_paragraph('2. Lõi trích xuất dữ liệu (OCR & LLM Pipeline):', style='List Bullet')
p1_2 = doc.add_paragraph()
p1_2.add_run('Đã có cấu trúc thư mục (ocr/, llm_predictions/) và kịch bản (orchestrator.py) để chạy quy trình đọc ảnh hóa đơn, sau đó dùng AI bóc tách thông tin thô thành định dạng JSON.')

doc.add_paragraph('3. Giao diện quản lý hóa đơn (Dashboard UI):', style='List Bullet')
p1_3 = doc.add_paragraph()
p1_3.add_run('Đã có giao diện Frontend (React) cho phép hiển thị hóa đơn "Chờ duyệt" và "Đã duyệt", hỗ trợ tính năng chọn file tải lên từ máy tính.')

doc.add_paragraph('4. Tính năng Xuất dữ liệu đa định dạng (Export):', style='List Bullet')
p1_4 = doc.add_paragraph()
p1_4.add_run('Hệ thống backend đã có các API xuất file chuẩn xác tương thích với phần mềm kế toán: File Excel (.csv), File định dạng MISA (.xml) và FAST (.csv).')

doc.add_paragraph('5. Đánh giá độ chính xác (Metrics & Ground Truth):', style='List Bullet')
p1_5 = doc.add_paragraph()
p1_5.add_run('Đã có công cụ chấm điểm AI, đối chiếu kết quả AI đọc được với đáp án chuẩn (Ground truth) để tính toán độ chính xác tỷ lệ lỗi.')

doc.add_paragraph('6. Cơ chế giới hạn (Freemium Quota):', style='List Bullet')
p1_6 = doc.add_paragraph()
p1_6.add_run('Đã tích hợp cơ chế chặn và cảnh báo khi người dùng vượt quá 200 hóa đơn miễn phí.')

# Phần 2: CÁC TÍNH NĂNG CẦN PHÁT TRIỂN
add_heading(doc, 'PHẦN 2: LỘ TRÌNH CÁC TÍNH NĂNG CẦN PHÁT TRIỂN (TƯƠNG LAI)')
p2 = doc.add_paragraph('Để biến dự án thành một phần mềm SaaS hoàn chỉnh, tiện dụng cho Kế toán, cần code thêm các tính năng sau:\n')

add_heading(doc, 'Giai đoạn 1: Tự động hóa nguồn vào (Data Ingestion)', level=2)
doc.add_paragraph('Auto Email Fetcher:', style='List Bullet')
doc.add_paragraph('Tự động kết nối vào Gmail/Outlook của công ty, cứ có email chứa file PDF hóa đơn là tự động tải về và đưa vào hệ thống chờ duyệt. (Giúp kế toán không phải tải bằng tay).')
doc.add_paragraph('Telegram/Zalo Bot Upload:', style='List Bullet')
doc.add_paragraph('Kế toán đi tiếp khách chỉ cần chụp ảnh bill gửi vào nhóm Zalo/Tele, hệ thống sẽ tự nhận diện và đẩy thẳng về Dashboard trên Web.')

add_heading(doc, 'Giai đoạn 2: Tối ưu Lõi Nghiệp vụ Kế Toán (Business Logic)', level=2)
doc.add_paragraph('Auto-Mapping (Ghép mã tự động):', style='List Bullet')
doc.add_paragraph('Tính năng cực kỳ quan trọng. Hệ thống tự động học thói quen người dùng: tự động hiểu "Máy tính T14" trên hóa đơn nhà cung cấp chính là mã kho "LAP01" trong hệ thống nội bộ để điền sẵn vào file Excel xuất ra.')
doc.add_paragraph('Toán học kiểm tra chéo (Cross-Validation):', style='List Bullet')
doc.add_paragraph('Viết thuật toán tự động cộng (Tổng tiền hàng + Thuế). Nếu không khớp với (Tổng thanh toán) do AI bị đọc nhầm một chữ số nào đó, hệ thống sẽ báo ĐỎ màn hình ngay lập tức để kế toán kịp thời kiểm tra.')

add_heading(doc, 'Giai đoạn 3: Rủi ro & Bảo mật pháp lý (Risk Management)', level=2)
doc.add_paragraph('Kiểm tra trạng thái Mã số thuế:', style='List Bullet')
doc.add_paragraph('Tích hợp API tra cứu mã số thuế. Nếu hóa đơn được xuất từ doanh nghiệp "Đã ngừng hoạt động/Bỏ trốn", hóa đơn sẽ bị đánh dấu đỏ nguy hiểm.')
doc.add_paragraph('Tài khoản Multi-Tenant (Dành cho kế toán dịch vụ):', style='List Bullet')
doc.add_paragraph('Chia Dashboard thành nhiều "Thư mục công ty" hoặc "Workspace" khác nhau để 1 người kế toán có thể quản lý số liệu cho 10-20 công ty cùng lúc mà không bị lẫn lộn hóa đơn.')

# Save doc
output_path = "Feature_List_Roadmap_AI_Invoice.docx"
doc.save(output_path)
print(f"Document saved to {os.path.abspath(output_path)}")
