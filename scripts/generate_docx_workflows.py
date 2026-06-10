import os
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 153)

doc = Document()

# Tiêu đề
title = doc.add_heading('MÔ TẢ CHI TIẾT LUỒNG HOẠT ĐỘNG & CÔNG NGHỆ\nCÁC TÍNH NĂNG ĐANG PHÁT TRIỂN', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('==================================================================').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('\nTài liệu này mô tả chi tiết cách thức vận hành thực tế của các tính năng sắp được lập trình, kèm theo CÔNG NGHỆ ĐỀ XUẤT (Tech Stack) cho Team Dev.\n')

# 1. Tính năng Auto Email Fetcher
add_heading(doc, '1. Tính năng Thu thập hóa đơn tự động qua Email (Auto Email Fetcher)')
p1 = doc.add_paragraph('Mục đích: Loại bỏ hoàn toàn thao tác Kế toán phải tự mở Email, tự lưu file PDF về máy tính.\n')
p1.add_run('Luồng hoạt động chi tiết:\n').bold = True
doc.add_paragraph('Bước 1: Setup (Cài đặt): Kế toán cung cấp cho hệ thống 1 địa chỉ email nhận hóa đơn của công ty (Ví dụ: hoadon@congty.com) và mật khẩu ứng dụng.', style='List Bullet')
doc.add_paragraph('Bước 2: Hoạt động ngầm (Cronjob): Một đoạn script Python chạy ngầm trên Server, cứ mỗi 5 phút sẽ tự động "gõ cửa" Gmail để kiểm tra hộp thư đến.', style='List Bullet')
doc.add_paragraph('Bước 3: Tải file: Nếu phát hiện có email mới chứa file đính kèm dạng .PDF, .PNG hoặc .XML, Code sẽ tự động tải file đó về thư mục "raw_invoices" trên máy chủ.', style='List Bullet')
doc.add_paragraph('Bước 4: Gọi AI xử lý: Sau khi tải xong, hệ thống đánh dấu Email đó là "Đã đọc" và đánh thức AI (orchestrator.py) dậy để trích xuất dữ liệu.', style='List Bullet')
doc.add_paragraph('Bước 5: Trả kết quả: Khi kế toán mở Dashboard Web lên, hóa đơn trong mail đã nằm sẵn ở trạng thái "Chờ duyệt".', style='List Bullet')
p1_tech = doc.add_paragraph()
p1_tech.add_run('🔧 CÔNG NGHỆ SỬ DỤNG:\n').bold = True
p1_tech.add_run('• Backend (Python): Thư viện `imaplib` để kết nối giao thức IMAP tải email, `email.message` để giải mã file đính kèm.\n')
p1_tech.add_run('• Tự động hóa: `APScheduler` hoặc `Celery` để hẹn giờ chạy ngầm 5 phút/lần.')

# 2. Tính năng Bot Telegram / Zalo
add_heading(doc, '2. Tính năng Gửi hóa đơn qua Bot Telegram / Zalo')
p2 = doc.add_paragraph('Mục đích: Tiện lợi cho nhân viên đi công tác, tiếp khách gửi hóa đơn trực tiếp bằng điện thoại.\n')
p2.add_run('Luồng hoạt động chi tiết:\n').bold = True
doc.add_paragraph('Bước 1: Chụp ảnh: Nhân viên cầm biên lai/hóa đơn, mở Telegram chat với con Bot của công ty (@AI_Invoice_Bot), chụp ảnh gửi vào chat.', style='List Bullet')
doc.add_paragraph('Bước 2: API Nhận lệnh: Máy chủ của Telegram nhận được ảnh, sẽ bắn 1 thông báo (Webhook) về máy chủ của chúng ta.', style='List Bullet')
doc.add_paragraph('Bước 3: Tải & Chuyển tiếp: Máy chủ của ta nhận được thông báo, tự động tải bức ảnh đó từ Telegram về, ném vào hàng đợi của Lõi AI.', style='List Bullet')
doc.add_paragraph('Bước 4: Trích xuất & Phản hồi: Lõi AI đọc xong, Bot Telegram sẽ nhắn lại cho nhân viên: "Đã nhận hóa đơn thành công. Tổng tiền: X".', style='List Bullet')
doc.add_paragraph('Bước 5: Đồng bộ Web: Bức ảnh và dữ liệu lập tức xuất hiện trên màn hình Dashboard.', style='List Bullet')
p2_tech = doc.add_paragraph()
p2_tech.add_run('🔧 CÔNG NGHỆ SỬ DỤNG:\n').bold = True
p2_tech.add_run('• Cổng giao tiếp: Thư viện `python-telegram-bot` (cho Telegram) hoặc `Zalo ZNS/OA API` (cho Zalo).\n')
p2_tech.add_run('• Backend API: Webhook Endpoint viết bằng `FastAPI` để lắng nghe tín hiệu từ máy chủ Telegram gửi về.\n')
p2_tech.add_run('• Tải file: `httpx` hoặc `requests` để tải ảnh về máy chủ gốc.')

# 3. Tính năng Auto-Mapping
add_heading(doc, '3. Tính năng AI Tự động Ghép Mã Hàng (Auto-Mapping)')
p3 = doc.add_paragraph('Mục đích: Giúp kế toán không phải chọn lại mã kho nội bộ cho từng dòng hàng hóa.\n')
p3.add_run('Luồng hoạt động chi tiết:\n').bold = True
doc.add_paragraph('Bước 1: Ghi nhớ lịch sử: Khi kế toán xử lý hóa đơn đầu tiên, AI đọc được tên hàng là "Giấy in A4 DoubleA". Kế toán sửa tên đó vào mã kho nội bộ là "VPP-001" và bấm Lưu.', style='List Bullet')
doc.add_paragraph('Bước 2: Cập nhật từ điển: Hệ thống Backend ngầm lưu lại một cặp giá trị vào Database: ["Giấy in A4 DoubleA" <=> "VPP-001"].', style='List Bullet')
doc.add_paragraph('Bước 3: Xử lý hóa đơn sau: Tháng sau, có một hóa đơn tương tự bay về. AI nhận diện được dòng chữ "Giấy in A4 Double A 70gsm".', style='List Bullet')
doc.add_paragraph('Bước 4: So sánh thông minh: Backend dùng thuật toán đối chiếu chuỗi, thấy dòng chữ này giống 85% so với từ khóa trong Database. Nó tự động điền luôn mã "VPP-001".', style='List Bullet')
doc.add_paragraph('Bước 5: Tiết kiệm thời gian: Kế toán nhìn thấy ô Mã Hàng đã được điền sẵn chính xác, chỉ việc bấm nút "Duyệt".', style='List Bullet')
p3_tech = doc.add_paragraph()
p3_tech.add_run('🔧 CÔNG NGHỆ SỬ DỤNG:\n').bold = True
p3_tech.add_run('• So sánh chuỗi truyền thống: Thư viện `RapidFuzz` (thuật toán Levenshtein distance) chuyên để tính toán phần trăm độ giống nhau của 2 đoạn text.\n')
p3_tech.add_run('• So sánh ngữ nghĩa (Nâng cao): Có thể dùng Vector Database (`ChromaDB` / `FAISS`) để lưu trữ từ vựng, kết hợp Embedding Model biến chữ thành vector để AI hiểu "Macbook" và "Máy tính Apple" là cùng 1 mã kho.')

# 4. Tính năng Toán học kiểm tra chéo
add_heading(doc, '4. Tính năng Kiểm Tra Chéo (Cross-Validation Logic)')
p4 = doc.add_paragraph('Mục đích: Phát hiện và cảnh báo sớm nếu AI nhận diện sai các con số nhạy cảm.\n')
p4.add_run('Luồng hoạt động chi tiết:\n').bold = True
doc.add_paragraph('Bước 1: Trích xuất: AI xuất ra: Tiền hàng chưa thuế = 1.000.000, Tiền thuế = 100.000, Tổng thanh toán = 1.700.000 (do AI nhìn nhầm số 1 thành 7).', style='List Bullet')
doc.add_paragraph('Bước 2: Tính toán ngầm: Backend chạy phép toán: 1.000.000 + 100.000 = 1.100.000.', style='List Bullet')
doc.add_paragraph('Bước 3: Phát hiện lỗi: Backend so sánh 1.100.000 với 1.700.000. Thấy không khớp nhau!', style='List Bullet')
doc.add_paragraph('Bước 4: Cảnh báo: Trên giao diện Web Dashboard, ô "Tổng thanh toán" lập tức nhấp nháy ĐỎ chót kèm cảnh báo.', style='List Bullet')
doc.add_paragraph('Bước 5: Sửa chữa: Kế toán sửa lại thành 1.100.000, ô báo lỗi mất đi, bấm Duyệt.', style='List Bullet')
p4_tech = doc.add_paragraph()
p4_tech.add_run('🔧 CÔNG NGHỆ SỬ DỤNG:\n').bold = True
p4_tech.add_run('• Backend: Logic kiểm tra `Pydantic validator` trong FastAPI để bắt lỗi toán học trước khi cho lưu vào DB.\n')
p4_tech.add_run('• Frontend: Quản lý trạng thái bằng `React State`, kết hợp `CSS Keyframes` để tạo hiệu ứng nhấp nháy viền ô input cảnh báo ĐỎ.')

# Save doc
output_path = "Chi_Tiet_Luong_Hoat_Dong.docx"
doc.save(output_path)
print(f"Document saved to {os.path.abspath(output_path)}")
