import os
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 153)

doc = Document()

# Tieu de
title = doc.add_heading('ĐỀ ÁN KINH DOANH & PHÁT TRIỂN SẢN PHẨM\nHỆ THỐNG AI INVOICE (XỬ LÝ HÓA ĐƠN TỰ ĐỘNG)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('==================================================================').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('\n')

# 1. Bài toán kế toán hiện tại
add_heading(doc, '1. Bài Toán Kế Toán Hiện Tại (Pain Points)')
p1 = doc.add_paragraph()
p1.add_run('Khối lượng công việc khổng lồ: ').bold = True
p1.add_run('Hàng tháng, phòng kế toán của các doanh nghiệp vừa và nhỏ (SME) phải tiếp nhận từ hàng trăm đến hàng ngàn hóa đơn đầu vào từ các nhà cung cấp khác nhau.\n')
p1.add_run('Nhập liệu thủ công dễ sai sót: ').bold = True
p1.add_run('Việc nhìn bằng mắt các file PDF/Ảnh hóa đơn sau đó gõ lại thủ công (Data Entry) vào phần mềm kế toán cực kỳ tốn thời gian, nhàm chán và tỷ lệ gõ nhầm (sai số tiền, sai mã số thuế) là rất cao.\n')
p1.add_run('Quy định khắt khe về thuế: ').bold = True
p1.add_run('Cơ quan thuế yêu cầu hóa đơn phải hoàn toàn hợp lệ. Nếu kế toán vô tình nhập hóa đơn của một "doanh nghiệp bỏ trốn", công ty sẽ bị phạt rất nặng.')

# 2. Họ đang dùng như thế nào
add_heading(doc, '2. Thực Trạng Xử Lý Của Kế Toán Hiện Nay')
p2 = doc.add_paragraph()
p2.add_run('Mức độ 1 - Thủ công hoàn toàn: ').bold = True
p2.add_run('Kế toán nhận email, in hóa đơn giấy hoặc mở file PDF trên máy tính, sau đó mở phần mềm (MISA, FAST) lên và gõ lại từng dòng thông tin mặt hàng, số tiền.\n')
p2.add_run('Mức độ 2 - Bán tự động (Dùng XML): ').bold = True
p2.add_run('Kế toán tải file XML của hóa đơn từ Email, sau đó dùng tính năng "Import XML" của phần mềm kế toán. \n')
p2.add_run('Tuy nhiên, nhược điểm chí mạng là: ').bold = True
p2.add_run('Tên mặt hàng của Nhà cung cấp (Vd: "Máy tính xách tay T14") không khớp với Mã nội bộ của Công ty (Vd: "LAP-01"). Kế toán vẫn phải ngồi "ghép mã" (mapping) thủ công cho từng dòng mặt hàng, vô cùng tốn thời gian.')

# 3. Web này sinh ra để giải quyết bài toán gì
add_heading(doc, '3. Hệ Thống AI Invoice Sinh Ra Để Giải Quyết Việc Gì?')
p3 = doc.add_paragraph()
p3.add_run('Dự án Web này hoạt động như một "Trợ lý Kế toán Ảo", giúp tự động hóa 90% quy trình trên:\n\n')
doc.add_paragraph('• Thu thập tự động: Tự động tải hóa đơn từ Email thay vì kế toán phải tải tay.', style='List Bullet')
doc.add_paragraph('• Trích xuất thông minh bằng AI (OCR & LLM): Đọc được cả hóa đơn PDF, ảnh chụp, bill giấy (những thứ mà XML không làm được), bóc tách chính xác từng trường thông tin.', style='List Bullet')
doc.add_paragraph('• Auto-Mapping (Ghép mã tự động): AI học từ lịch sử để tự động nhận diện "Máy tính xách tay" chính là mã "LAP-01", kế toán không cần ghép tay nữa.', style='List Bullet')
doc.add_paragraph('• Xuất chuẩn định dạng: Trả ra file Excel, MISA XML, FAST CSV để nạp thẳng vào phần mềm kế toán chỉ với 1 cú click.', style='List Bullet')

# 4. Cách bán Web này
add_heading(doc, '4. Chiến Lược Kinh Doanh (Tập trung SME & Kế toán dịch vụ)')
doc.add_paragraph('Hướng 1: Bán gói SaaS trả trước cho Doanh nghiệp Vừa và Nhỏ (SME)', style='List Bullet')
p4_1 = doc.add_paragraph('SME thường không có đủ ngân sách mua các hệ thống ERP khổng lồ nhưng lại rất cần một công cụ gọn nhẹ để giảm tải cho kế toán. Bán theo gói Quota số lượng hóa đơn xử lý:\n')
p4_1.add_run('- Gói Trải nghiệm (Phễu thu hút): Miễn phí 50 hóa đơn/tháng.\n')
p4_1.add_run('- Gói Cơ bản (SME): 300.000 VNĐ / tháng (xử lý tối đa 500 hóa đơn).\n')
p4_1.add_run('- Gói Chuyên nghiệp: 1.000.000 VNĐ / tháng (xử lý 2000 hóa đơn, kèm check Mã số thuế đỏ).')

doc.add_paragraph('Hướng 2: Bán tài khoản Multi-tenant cho "Kế toán dịch vụ/Đại lý thuế"', style='List Bullet')
p4_2 = doc.add_paragraph('Đánh thẳng vào tệp khách hàng là các cá nhân nhận làm sổ sách kế toán ngoài giờ hoặc các đại lý thuế. Một người kế toán dịch vụ thường làm cho 10-20 công ty nhỏ. Web sẽ cung cấp tính năng "Quản lý nhiều công ty" trên cùng 1 tài khoản. Bán theo dạng mua sỉ block (Ví dụ: Gói 10.000 hóa đơn dùng trong 1 năm, xài cho công ty nào cũng được).')

doc.add_paragraph('Hướng 3: Cung cấp API (White-label) cho các ứng dụng nhỏ', style='List Bullet')
p4_3 = doc.add_paragraph('Bán "lõi AI" dưới dạng API cho các phần mềm Quản lý bán hàng (POS), Quản lý phòng khám, Quản lý kho quy mô nhỏ. Các app này tự gọi API của hệ thống để quét hóa đơn, và hệ thống sẽ thu tiền theo từng lượt gọi (Pay-as-you-go).')

# 5. Những tính năng ăn tiền
add_heading(doc, '5. Các Tính Năng "Ăn Tiền" (Killer Features)')
p5 = doc.add_paragraph()
p5.add_run('1. Tính năng AI Auto-Mapping (Ghép mã tự động):\n').bold = True
p5.add_run('Đây là vũ khí quan trọng nhất. Phá vỡ rào cản của file XML truyền thống. Kế toán rất sợ việc phải tạo mới mã hàng hóa liên tục hoặc ghép mã tay. AI tự động học thói quen ghép mã của người dùng là tính năng khiến họ sẵn sàng rút hầu bao.\n\n')

p5.add_run('2. Xử lý "Chứng Từ Phi Tiêu Chuẩn":\n').bold = True
p5.add_run('Trong khi đối thủ (Bizzi, MISA) chỉ nhắm vào Hóa đơn điện tử chuẩn, hệ thống này "nhai" được cả Phiếu thu chi nội bộ, Biên lai viết tay, Hóa đơn nhiệt siêu thị mờ chữ nhờ công nghệ LLM và OCR mạnh mẽ.\n\n')

p5.add_run('3. Hệ thống Cảnh Báo "Doanh Nghiệp Rủi Ro" (Red Flags):\n').bold = True
p5.add_run('Tích hợp API check Tổng Cục Thuế. Nếu quét thấy hóa đơn từ doanh nghiệp vừa bị đóng mã số thuế hoặc bỏ trốn, màn hình Dashboard lập tức báo ĐỎ còi hú. Tính năng này cứu doanh nghiệp khỏi những án phạt hàng trăm triệu đồng.\n\n')

p5.add_run('4. Xuất file 1-Click (1-Click Export):\n').bold = True
p5.add_run('Tạo ra file XML chuẩn MISA và CSV chuẩn FAST. Kế toán không cần thao tác dán cột Excel rườm rà, tải file về kéo thả vào MISA là hoàn tất quy trình nhập liệu.')

# Save doc
output_path = "Business_Plan_AI_Invoice.docx"
doc.save(output_path)
print(f"Document saved to {os.path.abspath(output_path)}")
