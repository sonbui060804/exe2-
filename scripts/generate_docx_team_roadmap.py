import os
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 153)

doc = Document()

# Tiêu đề
title = doc.add_heading('LỘ TRÌNH PHÁT TRIỂN CÁC TÍNH NĂNG\n(Dành cho Team Phát Triển)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('==================================================================').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('\n')

add_heading(doc, 'Giai đoạn 1: Tự động hóa nguồn vào (Data Ingestion)', level=1)
p1_1 = doc.add_paragraph('Auto Email Fetcher:', style='List Bullet')
p1_1.runs[0].bold = True
doc.add_paragraph('Tự động kết nối vào Gmail/Outlook của công ty, cứ có email chứa file PDF hóa đơn là tự động tải về và đưa vào hệ thống chờ duyệt. (Giúp kế toán không phải tải bằng tay).')
p1_2 = doc.add_paragraph('Telegram/Zalo Bot Upload:', style='List Bullet')
p1_2.runs[0].bold = True
doc.add_paragraph('Kế toán đi tiếp khách chỉ cần chụp ảnh bill gửi vào nhóm Zalo/Tele, hệ thống sẽ tự nhận diện và đẩy thẳng về Dashboard trên Web.')

add_heading(doc, 'Giai đoạn 2: Tối ưu Lõi Nghiệp vụ Kế Toán (Business Logic)', level=1)
p2_1 = doc.add_paragraph('Auto-Mapping (Ghép mã tự động):', style='List Bullet')
p2_1.runs[0].bold = True
doc.add_paragraph('Tính năng cực kỳ quan trọng. Hệ thống tự động học thói quen người dùng: tự động hiểu "Máy tính T14" trên hóa đơn nhà cung cấp chính là mã kho "LAP01" trong hệ thống nội bộ để điền sẵn vào file Excel xuất ra.')
p2_2 = doc.add_paragraph('Toán học kiểm tra chéo (Cross-Validation):', style='List Bullet')
p2_2.runs[0].bold = True
doc.add_paragraph('Viết thuật toán tự động cộng (Tổng tiền hàng + Thuế). Nếu không khớp với (Tổng thanh toán) do AI bị đọc nhầm một chữ số nào đó, hệ thống sẽ báo ĐỎ màn hình ngay lập tức để kế toán kịp thời kiểm tra.')

add_heading(doc, 'Giai đoạn 3: Rủi ro & Bảo mật pháp lý (Risk Management)', level=1)
p3_1 = doc.add_paragraph('Kiểm tra trạng thái Mã số thuế:', style='List Bullet')
p3_1.runs[0].bold = True
doc.add_paragraph('Tích hợp API tra cứu mã số thuế. Nếu hóa đơn được xuất từ doanh nghiệp "Đã ngừng hoạt động/Bỏ trốn", hóa đơn sẽ bị đánh dấu đỏ nguy hiểm.')
p3_2 = doc.add_paragraph('Tài khoản Multi-Tenant (Dành cho kế toán dịch vụ):', style='List Bullet')
p3_2.runs[0].bold = True
doc.add_paragraph('Chia Dashboard thành nhiều "Thư mục công ty" hoặc "Workspace" khác nhau để 1 người kế toán có thể quản lý số liệu cho 10-20 công ty cùng lúc mà không bị lẫn lộn hóa đơn.')

# Save doc
output_path = "Lich_Trinh_Phat_Trien_Team.docx"
doc.save(output_path)
print(f"Document saved to {os.path.abspath(output_path)}")
